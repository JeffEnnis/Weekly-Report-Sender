"""
Weekly Market Intelligence Report — automation script (Outlook/Graph edition)
-------------------------------------------------------------------------------
Runs the search described in the Blueprint (Section 4), builds a formatted
.docx report, appends new bids/contacts to leads_tracker.xlsx, and emails
the report through your real Outlook mailbox via Microsoft Graph.

Designed to run inside the provided GitHub Actions workflow
(.github/workflows/weekly-report.yml) every Wednesday at 8:00 AM — see
Setup_Guide.docx for the full walkthrough. It also runs fine locally/manually.

SETUP (one-time):
  1. pip install anthropic python-docx openpyxl msal requests
  2. Set the environment variables listed in the CONFIG section below.
  3. Test manually:  python3 weekly_report.py
"""

import os
import json
import datetime
import base64
import sys
import logging

import anthropic
import requests
import msal
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl

# ----------------------------- LOGGING ---------------------------------
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ----------------------------- CONFIG ---------------------------------
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")          # required
MODEL = "claude-sonnet-4-6"

SEND_EMAIL = True

# Microsoft Graph / Outlook — from your Azure App Registration (see Setup Guide)
AZURE_CLIENT_ID = os.environ.get("AZURE_CLIENT_ID")
AZURE_TENANT_ID = os.environ.get("AZURE_TENANT_ID")
AZURE_CLIENT_SECRET = os.environ.get("AZURE_CLIENT_SECRET")
SENDER_EMAIL = os.environ.get("SENDER_EMAIL")          # the Outlook mailbox sending the report
REPORT_EMAIL_TO = os.environ.get("REPORT_EMAIL_TO")    # where the report should land

TRACKER_PATH = os.environ.get("LEADS_TRACKER_PATH", "leads_tracker.xlsx")
OUTPUT_DIR = os.environ.get("REPORT_OUTPUT_DIR", ".")

# Edit this as your focus shifts — regions, trades, target companies.
SEARCH_BRIEF = """
Search BC Bid, MERX, BC Hydro, TransLink, Metro Vancouver, the BC
Environmental Assessment Office project registry, Journal of Commerce,
Northern Miner, Mining.com, Business in Vancouver, and BC government news
releases for the past 7 days plus any newly posted or upcoming
opportunities. Focus on: heavy civil, general construction, mining/mineral
exploration, and civil/structural engineering, in British Columbia.

Return exactly:
- 3 Future Opportunities (projects likely to tender in 1-6 months, not yet
  formally posted — the best signals are EAO filings, municipal capital
  plans, and mining company investor updates)
- 3 Current Opportunities (open bids/RFPs closing in the next 30-60 days)
- 3 Current News items relevant to bidding or client strategy this week
- 3 Key Connections (named individuals worth contacting this week — role,
  company, why they matter, and a suggested outreach angle)

For each item give: title, one-line summary, source name, source URL,
date, and a suggested next action.

Respond with ONLY valid JSON, no markdown fences, no preamble, in this
exact shape:
{
  "future_opportunities": [{"title": "", "summary": "", "source": "", "url": "", "date": "", "next_action": ""}],
  "current_opportunities": [{"title": "", "summary": "", "source": "", "url": "", "date": "", "next_action": ""}],
  "current_news": [{"title": "", "summary": "", "source": "", "url": "", "date": "", "next_action": ""}],
  "key_connections": [{"title": "", "summary": "", "source": "", "url": "", "date": "", "next_action": "", "name": "", "company": ""}]
}
"""


def validate_json_structure(data: dict) -> bool:
    """Validate that the returned data has the expected structure."""
    required_keys = ["future_opportunities", "current_opportunities", "current_news", "key_connections"]
    for key in required_keys:
        if key not in data:
            logger.error(f"Missing required key: {key}")
            return False
        if not isinstance(data[key], list):
            logger.error(f"Key '{key}' is not a list")
            return False
    return True


def extract_json_from_text(raw: str) -> str:
    """
    Robustly extract JSON from potentially malformed text.
    Handles markdown fences, preamble, and extra text.
    """
    # Remove markdown code fences
    raw = raw.replace("```json", "").replace("```", "").strip()
    
    # Trim to outermost braces
    start = raw.find("{")
    end = raw.rfind("}")
    
    if start == -1 or end == -1 or end <= start:
        logger.error("No valid JSON braces found in response")
        logger.debug(f"Raw text (first 500 chars): {raw[:500]}")
        return ""
    
    extracted = raw[start:end + 1].strip()
    
    if not extracted:
        logger.error("Extracted JSON is empty")
        return ""
    
    return extracted


def fetch_weekly_data() -> dict:
    """
    Fetch weekly market data using Claude with web search.
    Includes comprehensive error handling and validation.
    """
    try:
        logger.info("Initializing Anthropic client...")
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        messages = [{"role": "user", "content": SEARCH_BRIEF}]
        response = None
        
        logger.info("Sending request to Claude with web search...")
        
        # Web search can take several search rounds to finish. If the model
        # pauses mid-search (stop_reason "pause_turn") or is cut off by the
        # token limit, feed its partial turn back in and let it continue,
        # rather than treating an incomplete response as a final answer.
        for attempt in range(6):
            logger.info(f"API call attempt {attempt + 1}/6...")
            try:
                response = client.messages.create(
                    model=MODEL,
                    max_tokens=8000,
                    tools=[{"type": "web_search_20250305", "name": "web_search"}],
                    messages=messages,
                )
            except anthropic.APIError as e:
                logger.error(f"API error on attempt {attempt + 1}: {e}")
                if attempt == 5:  # Last attempt
                    raise
                continue
            
            logger.info(f"Response received. Stop reason: {response.stop_reason}")
            
            if response.stop_reason == "pause_turn":
                logger.info("Model paused for more search. Continuing...")
                messages.append({"role": "assistant", "content": response.content})
                continue
            break
        
        if not response:
            raise RuntimeError("Failed to get response from Claude after 6 attempts")
        
        # Extract text from response content
        text_blocks = [b.text for b in response.content if hasattr(b, "text") and b.type == "text"]
        
        if not text_blocks:
            logger.error("No text blocks in response. Full response:")
            logger.error(response.model_dump_json(indent=2))
            raise RuntimeError(
                f"No text content in Claude response (stop_reason={response.stop_reason}). "
                "This may indicate the API is not returning search results."
            )
        
        raw = "\n".join(text_blocks).strip()
        logger.info(f"Received response ({len(raw)} chars). Extracting JSON...")
        
        # Extract JSON robustly
        json_str = extract_json_from_text(raw)
        
        if not json_str:
            logger.error("Could not extract valid JSON from response")
            logger.error(f"Raw response (first 1000 chars):\n{raw[:1000]}")
            raise RuntimeError("Response does not contain valid JSON")
        
        logger.info("Parsing JSON...")
        try:
            data = json.loads(json_str)
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error: {e}")
            logger.error(f"JSON string (first 500 chars): {json_str[:500]}")
            raise RuntimeError(f"Failed to parse JSON: {e}")
        
        # Validate structure
        if not validate_json_structure(data):
            logger.error("Invalid JSON structure. Expected keys not found.")
            logger.error(f"Received keys: {list(data.keys())}")
            raise RuntimeError("JSON structure does not match expected format")
        
        logger.info("Successfully fetched and parsed weekly data")
        return data
        
    except Exception as e:
        logger.error(f"Error fetching weekly data: {e}", exc_info=True)
        raise


def build_docx(data: dict, week_of: str) -> str:
    """Build a formatted Word document with the weekly report."""
    try:
        logger.info("Creating Word document...")
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title = doc.add_heading("Weekly Market Intelligence Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(f"Week of {week_of}  |  BC Construction / Engineering / Mining / Heavy Civil")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sections = [
            ("future_opportunities", "3 Future Opportunities"),
            ("current_opportunities", "3 Current Opportunities"),
            ("current_news", "3 Current News"),
            ("key_connections", "3 Key Connections"),
        ]

        for key, label in sections:
            logger.info(f"Adding section: {label}")
            doc.add_heading(label, level=1)
            items = data.get(key, [])
            logger.info(f"  Found {len(items)} items")
            
            for i, item in enumerate(items):
                try:
                    p = doc.add_paragraph()
                    run = p.add_run(item.get("title", "Untitled"))
                    run.bold = True
                    run.font.size = Pt(12)
                    doc.add_paragraph(item.get("summary", ""))
                    meta = doc.add_paragraph()
                    meta_run = meta.add_run(
                        f"Source: {item.get('source','')}  |  Date: {item.get('date','')}  |  "
                        f"Next action: {item.get('next_action','')}"
                    )
                    meta_run.italic = True
                    meta_run.font.size = Pt(9)
                    meta_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
                    if item.get("url"):
                        doc.add_paragraph(item["url"])
                    doc.add_paragraph()
                except Exception as e:
                    logger.warning(f"Error adding item {i} in section {key}: {e}")
                    continue

        filename = os.path.join(OUTPUT_DIR, f"Weekly_Report_{week_of}.docx")
        doc.save(filename)
        logger.info(f"Document saved to {filename}")
        return filename
        
    except Exception as e:
        logger.error(f"Error building DOCX: {e}", exc_info=True)
        raise


def append_to_tracker(data: dict, found_date: str):
    """Append new opportunities and contacts to the leads tracker spreadsheet."""
    try:
        if not os.path.exists(TRACKER_PATH):
            logger.warning(f"Tracker not found at {TRACKER_PATH} — skipping append.")
            return
        
        logger.info(f"Loading tracker from {TRACKER_PATH}...")
        wb = openpyxl.load_workbook(TRACKER_PATH)

        opp_ws = wb["Opportunities"]
        logger.info("Adding opportunities to tracker...")
        for stage_key, stage_label in [("future_opportunities", "Future"), ("current_opportunities", "Current")]:
            items = data.get(stage_key, [])
            logger.info(f"  Adding {len(items)} {stage_label} opportunities")
            for item in items:
                try:
                    opp_ws.append([
                        found_date, stage_label, item.get("title", ""), "", "", "", "",
                        "", item.get("source", ""), item.get("url", ""), "Not started",
                        item.get("next_action", ""), item.get("summary", ""),
                    ])
                except Exception as e:
                    logger.warning(f"Error appending opportunity: {e}")

        contacts_ws = wb["Contacts"]
        connections = data.get("key_connections", [])
        logger.info(f"Adding {len(connections)} contacts to tracker...")
        for item in connections:
            try:
                contacts_ws.append([
                    found_date, item.get("name", item.get("title", "")), "", item.get("company", ""),
                    "", item.get("source", ""), item.get("url", ""), "", "", "Not contacted",
                    item.get("next_action", ""), item.get("summary", ""),
                ])
            except Exception as e:
                logger.warning(f"Error appending contact: {e}")

        wb.save(TRACKER_PATH)
        logger.info("Tracker updated successfully")
        
    except Exception as e:
        logger.error(f"Error updating tracker: {e}", exc_info=True)
        raise


def get_graph_token() -> str:
    """Client-credentials auth against Azure AD, scoped to Microsoft Graph."""
    try:
        logger.info("Acquiring Microsoft Graph token...")
        logger.info(f"Using Azure Tenant: {AZURE_TENANT_ID}")
        logger.info(f"Using Azure Client ID: {AZURE_CLIENT_ID}")
        
        authority = f"https://login.microsoftonline.com/{AZURE_TENANT_ID}"
        app = msal.ConfidentialClientApplication(
            AZURE_CLIENT_ID, authority=authority, client_credential=AZURE_CLIENT_SECRET
        )
        result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
        
        # Debug: Log the full result structure
        logger.debug(f"MSAL result keys: {list(result.keys()) if isinstance(result, dict) else 'Not a dict'}")
        
        if "access_token" not in result:
            error_msg = result.get('error', 'Unknown error')
            error_description = result.get('error_description', '')
            full_error = f"{error_msg}: {error_description}" if error_description else error_msg
            
            logger.error(f"Failed to get Graph token")
            logger.error(f"  Error: {error_msg}")
            logger.error(f"  Description: {error_description}")
            logger.error(f"  Full result: {result}")
            
            raise RuntimeError(
                f"Authentication failed. Verify your AZURE_CLIENT_ID, AZURE_TENANT_ID, and AZURE_CLIENT_SECRET are correct. "
                f"Error: {full_error}"
            )
        
        logger.info("Graph token acquired successfully")
        return result["access_token"]
        
    except Exception as e:
        logger.error(f"Error getting Graph token: {e}", exc_info=True)
        raise


def send_email_via_graph(filepath: str, week_of: str):
    """Send the report email via Microsoft Graph API."""
    try:
        if not (AZURE_CLIENT_ID and AZURE_TENANT_ID and AZURE_CLIENT_SECRET and SENDER_EMAIL and REPORT_EMAIL_TO):
            logger.warning("Microsoft Graph credentials not fully set — skipping send. File saved at: " + filepath)
            return

        logger.info("Sending email via Outlook...")
        token = get_graph_token()
        
        with open(filepath, "rb") as f:
            content_b64 = base64.b64encode(f.read()).decode("utf-8")

        message = {
            "message": {
                "subject": f"Weekly Market Intelligence Report — {week_of}",
                "body": {"contentType": "Text", "content": "Attached: this week's bid, news, and connections report."},
                "toRecipients": [{"emailAddress": {"address": REPORT_EMAIL_TO}}],
                "attachments": [{
                    "@odata.type": "#microsoft.graph.fileAttachment",
                    "name": os.path.basename(filepath),
                    "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "contentBytes": content_b64,
                }],
            },
            "saveToSentItems": "true",
        }

        url = f"https://graph.microsoft.com/v1.0/users/{SENDER_EMAIL}/sendMail"
        
        logger.info(f"Posting to Microsoft Graph: {url}")
        resp = requests.post(
            url,
            headers={"Authorization": f"Bearer {token}"},
            json=message,
            timeout=30
        )
        
        if resp.status_code >= 300:
            logger.error(f"Graph sendMail failed ({resp.status_code}): {resp.text}")
            raise RuntimeError(f"Graph sendMail failed ({resp.status_code}): {resp.text}")
        
        logger.info("Email sent via Outlook successfully")
        
    except Exception as e:
        logger.error(f"Error sending email: {e}", exc_info=True)
        raise


def main():
    """Main entry point for the weekly report script."""
    try:
        logger.info("=" * 60)
        logger.info("Starting Weekly Market Intelligence Report")
        logger.info("=" * 60)
        
        if not ANTHROPIC_API_KEY:
            raise SystemExit("Set the ANTHROPIC_API_KEY environment variable before running.")

        week_of = datetime.date.today().isoformat()
        logger.info(f"Report date: {week_of}")
        
        logger.info("Step 1/4: Fetching weekly data...")
        data = fetch_weekly_data()

        logger.info("Step 2/4: Building report document...")
        filepath = build_docx(data, week_of)

        logger.info("Step 3/4: Updating leads tracker...")
        append_to_tracker(data, week_of)

        if SEND_EMAIL:
            logger.info("Step 4/4: Emailing report via Outlook...")
            send_email_via_graph(filepath, week_of)
        else:
            logger.info("Step 4/4: Email sending disabled")

        logger.info("=" * 60)
        logger.info(f"✓ Done. Report saved to {filepath}")
        logger.info("=" * 60)
        
    except Exception as e:
        logger.error("=" * 60)
        logger.error("✗ Script failed with error:")
        logger.error("=" * 60)
        logger.error(str(e), exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()
